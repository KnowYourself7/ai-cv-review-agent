from __future__ import annotations

from pathlib import Path
import re

from cv_review_agent.schemas import ParseResult


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def parse_resume_file(path: Path) -> ParseResult:
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return ParseResult(ok=False, error=f"Unsupported file type: {suffix}")
    if not path.exists() or path.stat().st_size == 0:
        return ParseResult(ok=False, error="Resume file is empty")

    try:
        if suffix == ".pdf":
            text = _parse_pdf(path)
        else:
            text = _parse_docx(path)
    except Exception as exc:
        return ParseResult(ok=False, error=f"Could not parse resume: {exc}")

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not normalized:
        return ParseResult(ok=False, error="No extractable text found. OCR is not supported in v1.")
    if not _has_readable_resume_text(normalized):
        return ParseResult(
            ok=False,
            error="No readable resume text found. Upload a text-based PDF or DOCX; scanned or encoded PDFs need OCR.",
        )
    return ParseResult(ok=True, text=normalized)


def _parse_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _parse_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _has_readable_resume_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    if len(compact) < 80:
        return False

    repeated_lines = text.splitlines()
    if repeated_lines:
        most_common_count = max(repeated_lines.count(line) for line in set(repeated_lines))
        if most_common_count / len(repeated_lines) > 0.6 and len(set(repeated_lines)) <= 3:
            return False

    readable_chars = re.findall(r"[\u4e00-\u9fffA-Za-z]", compact)
    if len(readable_chars) / len(compact) < 0.45:
        return False

    resume_signals = [
        "resume",
        "cv",
        "education",
        "experience",
        "skills",
        "project",
        "email",
        "phone",
        "简历",
        "教育",
        "学历",
        "经历",
        "经验",
        "技能",
        "项目",
        "邮箱",
        "电话",
    ]
    lowered = text.lower()
    return any(signal in lowered for signal in resume_signals)
